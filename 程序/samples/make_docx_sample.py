"""生成测试用 Word 文档（含标题层级、表格、页眉页脚、加粗与项目符号）。

用法（在 程序/ 目录下）：python samples/make_docx_sample.py
输出：samples/sample_doc.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent / "sample_doc.docx"


def build():
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "Journal of Educational Robotics"
    doc.sections[0].footer.paragraphs[0].text = "Preprint version, 2026"

    doc.add_heading("Observing Robot Failures in Classrooms", level=0)
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "According to productive failure theory, experiencing failure during "
        "problem-solving can enhance knowledge acquisition in subsequent "
        "instruction. We designed a social robot-assisted teaching activity "
        "in which students observe unsuccessful problem-solving attempts.")

    doc.add_heading("Methods", level=1)
    p = doc.add_paragraph("We compared three instructional conditions. ")
    r = p.add_run("The robot failure condition")
    r.bold = True
    p.add_run(" produced the largest conceptual gain.")

    doc.add_paragraph("Random assignment to conditions", style="List Bullet")
    doc.add_paragraph("Pretest and posttest measures", style="List Bullet")

    doc.add_heading("Results", level=1)
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    data = [["Condition", "Participants", "Outcome"],
            ["Robot failure", "46 students", "High improvement"],
            ["Direct instruction", "44 students", "Small improvement"]]
    for row, values in zip(table.rows, data):
        for cell, v in zip(row.cells, values):
            cell.text = v

    doc.add_paragraph()
    doc.add_paragraph("Numbers such as 135 and symbols like — are left alone.")
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.size is None:
                r.font.size = Pt(11)
    doc.save(str(OUT))
    print(f"已生成：{OUT}")


if __name__ == "__main__":
    build()
